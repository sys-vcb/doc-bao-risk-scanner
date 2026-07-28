import os
from datetime import datetime
from typing import List
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.config import REPORTS_DIR
from app.models import RiskItem

def set_cell_background(cell, fill_hex: str):
    """Đặt màu nền cho cell trong bảng docx"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def add_hyperlink(paragraph, url: str, text: str, color_hex="0066CC"):
    """Thêm đường link bài gốc vào paragraph trong docx"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def generate_daily_docx_report(risk_items: List[RiskItem], target_date: str = None) -> str:
    """
    Tổng hợp danh sách tin rủi ro thành 1 file Word (.docx) chuẩn 5 cột theo quy định.
    Tên file: Bao_Cao_Rui_Ro_YYYY-MM-DD.docx
    """
    if not target_date:
        target_date = datetime.now().strftime("%Y-%m-%d")

    filename = f"Bao_Cao_Rui_Ro_{target_date}.docx"
    file_path = REPORTS_DIR / filename

    doc = Document()

    # Cấu hình Margins (Lề trang)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 1. Header Báo cáo
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO CẢNH BÁO RỦI RO DOANH NGHIỆP & TỔ CHỨC CÁ NHÂN")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42) # Slate Dark

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Ngày tổng hợp: {target_date}  |  Tổng số phát hiện: {len(risk_items)} rủi ro")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph() # Khoảng trống

    # 2. Bảng Dữ Liệu 5 Cột
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Tiêu đề 5 Cột chuẩn
    headers = [
        "1. Tên Cá nhân/ Doanh nghiệp/ Tổ chức",
        "2. Nội dung tóm tắt rủi ro",
        "3. Khu vực",
        "4. Loại rủi ro / Từ khóa",
        "5. Ngày & Link bài gốc"
    ]
    
    # Độ rộng 5 cột (Tổng 6.9 inches)
    col_widths = [Inches(1.5), Inches(2.2), Inches(1.0), Inches(1.1), Inches(1.1)]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B") # Dark Indigo Header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].width = col_widths[i]

    # Điền dữ liệu các dòng
    if not risk_items:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[4])
        row_cells[0].text = "Không ghi nhận bài viết cảnh báo rủi ro nào trong ngày."
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
    else:
        for idx, item in enumerate(risk_items):
            row_cells = table.add_row().cells
            bg_color = "F8FAFC" if idx % 2 == 1 else "FFFFFF"

            # Cột 1: Tên Doanh nghiệp / Cá nhân
            row_cells[0].text = item.entity_name or "Chưa xác định"
            
            # Cột 2: Tóm tắt rủi ro
            row_cells[1].text = item.summary or ""
            
            # Cột 3: Khu vực
            row_cells[2].text = item.province or "Toàn quốc"
            
            # Cột 4: Loại rủi ro
            row_cells[3].text = item.risk_type or ""
            
            # Cột 5: Ngày & Link bài gốc
            p5 = row_cells[4].paragraphs[0]
            p5.text = f"{item.published_date}\n"
            if item.source_url:
                add_hyperlink(p5, item.source_url, "Xem bài gốc", color_hex="2563EB")

            # Định dạng chung cho các cells trong hàng
            for i in range(5):
                cell = row_cells[i]
                cell.width = col_widths[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_background(cell, bg_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(9.5)

    doc.save(str(file_path))
    return str(file_path)
